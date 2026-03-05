"""
export_utils.py — Document Export Utilities
--------------------------------------------
Export generated content to PDF and DOCX formats
"""

import os
import tempfile
from datetime import datetime
from pathlib import Path


def export_to_pdf(content: str, title: str = "Study Material") -> str:
    """Export content to PDF using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        raise ImportError("reportlab not installed. Run: pip install reportlab")
    
    # Create temp file
    temp_dir = tempfile.gettempdir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = os.path.join(temp_dir, f"{title}_{timestamp}.pdf")
    
    # Create PDF
    doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#2c3e50',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    # Add title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Add timestamp
    date_style = styles['Normal']
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", date_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Add content
    content_style = styles['BodyText']
    
    # Split content by paragraphs
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Handle markdown-style headers
            if para.startswith('# '):
                story.append(Paragraph(para[2:], styles['Heading1']))
            elif para.startswith('## '):
                story.append(Paragraph(para[3:], styles['Heading2']))
            elif para.startswith('### '):
                story.append(Paragraph(para[4:], styles['Heading3']))
            else:
                # Regular paragraph
                story.append(Paragraph(para.replace('\n', '<br/>'), content_style))
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    return pdf_path


def export_to_docx(content: str, title: str = "Study Material") -> str:
    """Export content to DOCX using python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    # Create temp file
    temp_dir = tempfile.gettempdir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(temp_dir, f"{title}_{timestamp}.docx")
    
    # Create document
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Add content
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Handle markdown-style headers
            if para.startswith('# '):
                doc.add_heading(para[2:], level=1)
            elif para.startswith('## '):
                doc.add_heading(para[3:], level=2)
            elif para.startswith('### '):
                doc.add_heading(para[4:], level=3)
            elif para.startswith('- ') or para.startswith('* '):
                # Bullet point
                doc.add_paragraph(para[2:], style='List Bullet')
            else:
                # Regular paragraph
                doc.add_paragraph(para)
    
    doc.save(docx_path)
    return docx_path
